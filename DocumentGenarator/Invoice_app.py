import streamlit as st
import pandas as pd
from openai import OpenAI
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from datetime import datetime
from reportlab.lib.utils import ImageReader

st.set_page_config(page_title="Invoice Summary Generator", layout="wide")
st.title("🧾 Smart Invoice Summary Generator (LLM-Powered)")

st.markdown("""
Upload your fee data (CSV) and get professional, plain-English invoice summaries generated using OpenAI's GPT.
Each row in your file should contain one client's invoice information.
""")

# API Key input
api_key = st.secrets["OPENAI_API_KEY"] if "OPENAI_API_KEY" in st.secrets else st.text_input("Enter OpenAI API Key", type="password")
client = OpenAI(api_key=api_key)

# File upload
uploaded_file = st.file_uploader("Upload Fee CSV", type="csv")

# Optional branding image path (update with your logo file path)
LOGO_PATH = "logo.png"

# Prompt template
prompt_template = """
Generate a plain-English invoice summary for the following client:

Client Name: {client_name}
Quarter: {quarter}
Valuation: ${valuation}
Management Fee: ${management_fee}
Performance Fee: ${performance_fee}

Explain what each fee means and provide a professional but friendly tone.
"""

def generate_invoice_document(row):
    try:
        prompt = prompt_template.format(**row)
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        return f" Error: {e}"

def generate_llm_summary(row):
    try:
        prompt = prompt_template.format(**row)
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0.7
        )
        return response.choices[0].message['content']
    except Exception as e:
        return f" Error: {e}"

def generate_docx(name, summary, invoice_id):
    doc = Document()
    doc.add_heading(f"Invoice Summary - {name} (#{invoice_id})", level=1)
    for line in summary.split('\n'):
        doc.add_paragraph(line.strip())
    doc.add_paragraph("\nThank you for your business.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf(name, summary, invoice_id):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    y = height - 50

    # Optional logo
    try:
        logo = ImageReader(LOGO_PATH)
        c.drawImage(logo, 50, y - 60, width=100, preserveAspectRatio=True)
    except:
        pass  # If logo not found, skip

    c.setFont("Helvetica-Bold", 14)
    c.drawString(200, y, f"Invoice Summary - {name} (#{invoice_id})")
    y -= 100

    c.setFont("Helvetica", 11)
    for line in summary.split("\n"):
        if y < 100:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 11)
        c.drawString(50, y, line.strip())
        y -= 20

    # Footer
    c.setFont("Helvetica-Oblique", 9)
    c.drawString(50, 30, "This document was automatically generated.")
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer


if uploaded_file:
    df = pd.read_csv(uploaded_file)

    required_columns = ["client_name", "valuation", "management_fee", "performance_fee", "quarter"]
    if all(col in df.columns for col in required_columns):
        if api_key:
            st.info("Generating summaries using OpenAI...")
            df["llm_summary"] = df.apply(generate_invoice_document, axis=1)
            st.success("Summaries generated below:")
            today = datetime.today().strftime("%Y%m%d")

            for i, row in df.iterrows():
                st.markdown("---")

                # Generate invoice ID from name + quarter + date
                today = datetime.today().strftime("%Y%m%d")
                shortname = row["client_name"].split()[0]
                invoice_id = f"{shortname}_{row['quarter'].replace('-', '')}_{today}"

                st.markdown(f"**{row['client_name']} ({row['quarter']}) - Invoice #{invoice_id}**")
                st.markdown(row["llm_summary"])

                docx_buffer = generate_docx(row["client_name"], row["llm_summary"], invoice_id)
                pdf_buffer = generate_pdf(row["client_name"], row["llm_summary"], invoice_id)
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="Download DOCX",
                        data=docx_buffer,
                        file_name=f"invoice_{invoice_id}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with col2:
                    st.download_button(
                        label="Download PDF",
                        data=pdf_buffer,
                        file_name=f"invoice_{invoice_id}.pdf",
                        mime="application/pdf"
                    )

        else:
            st.warning("Please enter your OpenAI API key to generate summaries.")
    else:
        st.error("CSV must include the following columns: " + ", ".join(required_columns))
